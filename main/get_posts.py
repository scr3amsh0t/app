import requests
from main.main import get_name, get_screen_name
import json
import csv
import docx
import logging
from django.conf import settings
import os

# logging.basicConfig(level=logging.DEBUG, format='%(asctime)s - %(levelname)s - %(message)s')
# logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
# logging.basicConfig(level=logging.WARNING, format='%(asctime)s - %(levelname)s - %(message)s')
# logging.basicConfig(level=logging.ERROR, format='%(asctime)s - %(levelname)s - %(message)s')
# logging.basicConfig(level=logging.CRITICAL, format='%(asctime)s - %(levelname)s - %(message)s')

def get_posts_count(screen_name):
    token = 'vk1.a.byMJTaFR8uzQ2VOgF72GpGczOd0RnOu1YBVklpdL9Rnndd-5TSH1FGz94XMiFgw4b13TFUQNikYHk79VQ5jwJ7GHKIoVZb3No7t97wJZTlgj5iqirPrXCXikDQOuSewYbYUbwuMb7kth4YqsAC8pDxBE-ax68I0qYiEHhkFnumJo3HzsWxRgvfPKwMck6jl1IDxVnpZ_uTGQMAZa2Kl9Xg'
    version = 5.199
    owner_id = screen_name
    count = 1
    response = requests.get('https://api.vk.com/method/wall.get',
                            params={
                                'access_token': token,
                                'v': version,
                                'owner_id': owner_id,
                                'count': count
                            })
    posts_count = response.json()['response']['count']
    return posts_count


def get_posts(screen_name, posts_count):
    token = 'vk1.a.byMJTaFR8uzQ2VOgF72GpGczOd0RnOu1YBVklpdL9Rnndd-5TSH1FGz94XMiFgw4b13TFUQNikYHk79VQ5jwJ7GHKIoVZb3No7t97wJZTlgj5iqirPrXCXikDQOuSewYbYUbwuMb7kth4YqsAC8pDxBE-ax68I0qYiEHhkFnumJo3HzsWxRgvfPKwMck6jl1IDxVnpZ_uTGQMAZa2Kl9Xg'
    version = 5.199
    owner_id = screen_name
    offset = 0
    posts = []
    if int(posts_count) > 100:
        count_count = 100
        while (offset+100) < posts_count:
            response = requests.get('https://api.vk.com/method/wall.get',
                                    params={
                                        'access_token': token,
                                        'v': version,
                                        'owner_id': owner_id,
                                        'count': count_count,
                                        'offset': offset
                                    })
            data = response.json()['response']['items']
            offset += 100
            posts.extend(data)
        count_count = posts_count - offset
        response = requests.get('https://api.vk.com/method/wall.get',
                                params={
                                    'access_token': token,
                                    'v': version,
                                    'owner_id': owner_id,
                                    'count': count_count,
                                    'offset': offset
                                })
        data = response.json()['response']['items']
        posts.extend(data)
    else:
        response = requests.get('https://api.vk.com/method/wall.get',
                                params={
                                    'access_token': token,
                                    'v': version,
                                    'owner_id': owner_id,
                                    'count': posts_count,
                                    'offset': offset
                                })
        data = response.json()['response']['items']
        posts.extend(data)
    return posts

def make_json(data):
    texts = []
    for post in data:
        texts.append(post['text'])
    texts = [line.replace("\n", "") for line in texts]
    data = []
    for text in texts:
        data.append({"text": text})
    return data


url_module_2 = 'http://localhost:8002/api'

def zapros(data):
    response = requests.post(url_module_2, data = json.dumps(data))
    if response.status_code == 202:
        token = response.json()['id']
    while (response.status_code != 200):
        response = requests.get(f'{url_module_2}/{token}')
        answers = response.json()
    if (response.status_code != 200 & response.status_code != 200): 
        print(response.status_code)
    return answers

def posts_txt(data):
    text = ''
    i = 1
    for answer in data:
        text += "Teкст поста " + str(i) + ":\n" + str(answer['text']) + "\nЗапрещенный контент найден?: " + str(answer['result']) + "\n"
        i+=1
    return text

def posts_txt_with_id(data, posts):
    text = []
    for i in range(len(data)):
        text.append("Teкст поста " + str(i+1) + ":\n" + str(data[i]['text']) + "\nЗапрещенный контент найден?: " + str(data[i]['result']) + "\n")
        text.append("Ссылка на пост:"+ " \nhttps://vk.com/wall" + str(posts[i]['from_id']) + "_" + str(posts[i]['id']) + "\n\n")        
    return text

def posts_csv(data):
    with open('posts.csv', 'w', newline='', encoding='utf-8') as f:
        csv_writer = csv.writer(f)
        csv_writer.writerow(['Номер поста', 'Текст поста', 'Запрещенный контент найден?'])
        i = 1
        for post in data:
            csv_writer.writerow([i, post['text'], post['result']])
            i+=1
    return 'posts.csv'

def posts_csv_with_id(data, posts):
    with open('posts_id.csv', 'w', newline='', encoding='utf-8') as f:
        csv_writer = csv.writer(f)
        csv_writer.writerow(['Номер поста', 'Текст поста', 'Запрещенный контент найден?', 'Ссылка на пост'])
        j = 1
        for i in range(len(data)):
            csv_writer.writerow([j, data[i]['text'], data[i]['result'], 'https://vk.com/wall' + str(posts[i]['from_id']) + '_' + str(posts[i]['id'])])
            j+=1
    return 'posts_id.csv'

def posts_docx(data):
    doc = docx.Document()
    i = 1
    for post in data:
        paragraph = doc.add_paragraph()
        paragraph.add_run("Teкст поста " + str(i) + ":\n" + post['text'] + "\nЗапрещенный контент найден?: " + str(post['result']))
        i+=1
        file_path = "posts.docx"
        doc.save(file_path)
    return file_path

def posts_docx_with_id(data, posts):
    doc = docx.Document()
    for i in range(len(data)):
        paragraph = doc.add_paragraph()
        paragraph.add_run("Teкст поста " + str(i+1) + ":\n" + str(data[i]['text']) + "\nЗапрещенный контент найден?: " + str(data[i]['result']) + "\n" + "Ссылка на пост:"+ " \nhttps://vk.com/wall" + str(posts[i]['from_id']) + "_" + str(posts[i]['id']) + "\n")
        file_path = "posts_id.docx"
        doc.save(file_path)
    return file_path