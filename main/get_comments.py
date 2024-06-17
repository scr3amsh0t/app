import time
import requests
from main.main import get_name, get_screen_name
import logging
import json
import csv
import docx

# logging.basicConfig(level=logging.DEBUG, format='%(asctime)s - %(levelname)s - %(message)s')
# logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
# logging.basicConfig(level=logging.WARNING, format='%(asctime)s - %(levelname)s - %(message)s')
# logging.basicConfig(level=logging.ERROR, format='%(asctime)s - %(levelname)s - %(message)s')
# logging.basicConfig(level=logging.CRITICAL, format='%(asctime)s - %(levelname)s - %(message)s')


def get_comments(screen_name,posts_count):
    token = 'vk1.a.byMJTaFR8uzQ2VOgF72GpGczOd0RnOu1YBVklpdL9Rnndd-5TSH1FGz94XMiFgw4b13TFUQNikYHk79VQ5jwJ7GHKIoVZb3No7t97wJZTlgj5iqirPrXCXikDQOuSewYbYUbwuMb7kth4YqsAC8pDxBE-ax68I0qYiEHhkFnumJo3HzsWxRgvfPKwMck6jl1IDxVnpZ_uTGQMAZa2Kl9Xg'
    version = 5.199
    domain = screen_name
    count = 1
    offset = 0
    comments = []
    thread_items_count = 10
    while offset < int(posts_count):
        response = requests.get('https://api.vk.com/method/wall.get',
                                params={
                                    'access_token': token,
                                    'v': version,
                                    'domain': domain,
                                    'offset': offset,
                                    'count': count
                                })
        offset += 1
        posts = response.json()['response']['items']
        time.sleep(0.5)
        for post in posts:
            posts_id = post['id']
            comments_count = post['comments']['count']
            response = requests.get('https://api.vk.com/method/wall.getComments',
                                    params={
                                        'access_token': token,
                                        'v': version,
                                        'owner_id': screen_name,
                                        'post_id': posts_id,
                                        'count': comments_count,
                                        "thread_items_count": thread_items_count
                                    })
            data = response.json()['response']['items']
            comments.extend(data)
    return comments

def make_json(data):
    texts = []
    for comment in data:
        if comment['text'] != '':
            texts.append(comment['text'])
        threads = comment['thread']['items']
        for thread in threads:
                if thread['text'] != '':
                    texts.append(thread['text'])
    texts = [line.replace("\n", "") for line in texts]
    data = []
    for text in texts:
        data.append({"text": text})
    return data

def make_id_json(data):
    ids = []
    for comment in data:
        ids.append(comment['from_id'])
        threads = comment['thread']['items']
        for thread in threads:
                ids.append(thread['from_id'])
    from_id = []
    for id in ids:
        from_id.append({"from_id" : id})
    return from_id

def make_comment_id_json(data):
    comm_ids = []
    for comment in data:
        comm_ids.append((comment['id'], comment['parents_stack']))
        threads = comment['thread']['items']
        for thread in threads:
                comm_ids.append((thread['id'], thread['parents_stack']))
    comm_id = []
    for id, parents in comm_ids:
        comm_id.append({"id" : id, "parents" : parents})
    return comm_id

url_module_2 = 'http://localhost:8002/api'

def zapros(data):
    response = requests.post(url_module_2, data = json.dumps(data))
    if response.status_code == 202:
        token = response.json()['id']
    while (response.status_code != 200):
        response = requests.get(f'{url_module_2}/{token}')
        answers = response.json()
    if (response.status_code != 200 & response.status_code != 200): 
        print(response.status_code)
    return answers
     
def comments_txt(data):
    text = ''
    i = 1
    for answer in data:
        text += "Teкст комментария " + str(i) + ":\n" + str(answer['text']) + "\nЗапрещенный контент найден?: " + str(answer['result']) + "\n"
        i+=1
    return text

def comments_txt_with_id(data, from_id):
    text = []
    for i in range(len(data)):
        text.append("Teкст комментария " + str(i+1) + ":\n" + str(data[i]['text']) + "\nЗапрещенный контент найден?: " + str(data[i]['result']) + "\n")
        text.append("Ссылка на автора:" + " \nhttps://vk.com/id" + str(from_id[i]['from_id']) + "\n\n")        
    return text

def comments_txt_with_comm_id(data, comms, comms_id):
    text = []
    for i in comms:
        a = ("Ссылка: " + "https://vk.com/wall" + str(i['owner_id']) + "_" + str(i['post_id']))
    for i in range(len(data)):
        text.append("Текст комментария " + str(i+1) + ":\n" + str(data[i]['text']) + "\nЗапрещенный контент найден?: " + str(data[i]['result']) + "\n")
        text.append(str(a) + "?reply=" + str(comms_id[i]['id']) + "\n\n")
    return text

def comments_csv(data):
    with open('comments.csv', 'w', newline='', encoding='utf-8') as f:
        csv_writer = csv.writer(f)
        csv_writer.writerow(['Номер комментария', 'Текст комментария', 'Запрещенный контент найден?'])
        i = 1
        for comment in data:
            csv_writer.writerow([i, comment['text'], comment['result']])
            i+=1
    return 'comments.csv'

def comments_csv_with_id(data, from_id):
        with open('comments_with_author_id.csv', 'w', newline='', encoding='utf-8') as f:
            csv_writer = csv.writer(f)
            csv_writer.writerow(['Номер комментария', 'Текст комментария', 'Запрещенный контент найден?', 'Ссылка на автора'])
            j = 1
            for i in range(len(data)):
                csv_writer.writerow([j, data[i]['text'], data[i]['result'], 'https://vk.com/id' + str(from_id[i]['from_id'])])
                j+=1
        return 'comments_with_author_id.csv'

def comments_docx(data):
    doc = docx.Document()
    i = 1
    for comment in data:
        paragraph = doc.add_paragraph()
        paragraph.add_run("Teкст комментария " + str(i) + ":\n" + comment['text'] + "\nЗапрещенный контент найден?: " + str(comment['result']))
        i+=1
        file_path = "comments.docx"
        doc.save(file_path)
    return file_path


def comments_docx_with_id(data, from_id):
    doc = docx.Document()
    for i in range(len(data)):
        paragraph = doc.add_paragraph()
        paragraph.add_run("Teкст комментария " + str(i+1) + ":\n" + str(data[i]['text']) + "\nЗапрещенный контент найден?: " + str(data[i]['result']) + "\n" + "Ссылка на автора:"+ " \nhttps://vk.com/id" + str(from_id[i]['from_id']) + "\n")
        file_path = "comments_with_author_id.docx"
        doc.save(file_path)
    return file_path

